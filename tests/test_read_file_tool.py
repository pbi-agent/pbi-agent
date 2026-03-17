from __future__ import annotations

import sys
from pathlib import Path
from types import ModuleType, SimpleNamespace

import pytest

from pbi_agent.tools import read_file as read_file_tool
from pbi_agent.tools.types import ToolContext


def test_read_file_returns_requested_line_window(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.chdir(tmp_path)
    (tmp_path / "notes.txt").write_text("one\ntwo\nthree\nfour\n", encoding="utf-8")

    result = read_file_tool.handle(
        {"path": "notes.txt", "start_line": 2, "max_lines": 2},
        ToolContext(),
    )

    assert result == {
        "path": "notes.txt",
        "start_line": 2,
        "end_line": 3,
        "total_lines": 4,
        "content": "two\nthree\n",
        "has_more_lines": True,
        "windowed": True,
    }


def test_read_file_auto_detects_utf16_bom(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.chdir(tmp_path)
    (tmp_path / "utf16.txt").write_bytes("hello\nworld\n".encode("utf-16"))

    result = read_file_tool.handle({"path": "utf16.txt"}, ToolContext())

    assert result["path"] == "utf16.txt"
    assert result["content"] == "hello\nworld\n"
    assert result["has_more_lines"] is False
    assert "windowed" not in result


def test_read_file_summarizes_csv_with_schema_and_stats(
    tmp_path: Path, monkeypatch
) -> None:
    pytest.importorskip("polars")
    monkeypatch.chdir(tmp_path)
    (tmp_path / "dataset.csv").write_text(
        "city,sales,ordered_at\nSeattle,10,2025-01-01\nSeattle,20,2025-01-02\nPortland,30,2025-01-03\n",
        encoding="utf-8",
    )

    result = read_file_tool.handle({"path": "dataset.csv"}, ToolContext())

    assert result["shape"] == {"rows": 3, "columns": 3}
    assert (
        result["summary"]
        == "3 rows x 3 columns; column mix: 1 numeric, 1 datetime, 1 text; missing values: no missing values."
    )
    assert "- city: String; kind=text; examples=Seattle, Portland" in result["schema"]
    assert "- sales: Int64; kind=numeric; min=10; max=30; mean=20" in result["schema"]
    assert (
        "- ordered_at: Date; kind=datetime; range=2025-01-01..2025-01-03"
        in result["schema"]
    )
    assert result["preview"] == (
        "city,sales,ordered_at\n"
        "Seattle,10,2025-01-01\n"
        "Seattle,20,2025-01-02\n"
        "Portland,30,2025-01-03\n"
    )


def test_read_file_summarizes_csv_with_cr_only_line_endings(
    tmp_path: Path, monkeypatch
) -> None:
    pytest.importorskip("polars")
    monkeypatch.chdir(tmp_path)
    (tmp_path / "dataset.csv").write_bytes(
        b"city,sales,ordered_at\rSeattle,10,2025-01-01\rPortland,30,2025-01-03\r"
    )

    result = read_file_tool.handle({"path": "dataset.csv"}, ToolContext())

    assert result["shape"] == {"rows": 2, "columns": 3}
    assert (
        result["summary"]
        == "2 rows x 3 columns; column mix: 1 numeric, 1 datetime, 1 text; missing values: no missing values."
    )
    assert (
        "- ordered_at: Date; kind=datetime; range=2025-01-01..2025-01-03"
        in result["schema"]
    )
    assert "- sales: Int64; kind=numeric; min=10; max=30; mean=20" in result["schema"]
    assert result["preview"] == (
        "city,sales,ordered_at\nSeattle,10,2025-01-01\nPortland,30,2025-01-03\n"
    )


def test_read_file_returns_all_excel_sheets(tmp_path: Path, monkeypatch) -> None:
    pl = pytest.importorskip("polars")
    monkeypatch.chdir(tmp_path)
    (tmp_path / "workbook.xlsx").write_bytes(b"placeholder")

    monkeypatch.setattr(
        read_file_tool,
        "_read_excel_workbook",
        lambda path: {
            "Orders": pl.DataFrame(
                {"city": ["Seattle", "Portland"], "sales": [10, 20]}
            ),
            "Returns": pl.DataFrame({"city": ["Seattle"], "count": [1]}),
        },
    )

    result = read_file_tool.handle({"path": "workbook.xlsx"}, ToolContext())

    assert result["path"] == "workbook.xlsx"
    assert result["sheet_count"] == 2
    assert [sheet["name"] for sheet in result["sheets"]] == ["Orders", "Returns"]
    assert result["sheets"][0]["shape"] == {"rows": 2, "columns": 2}
    assert result["sheets"][0]["preview"] == ("city,sales\nSeattle,10\nPortland,20\n")
    assert result["sheets"][1]["shape"] == {"rows": 1, "columns": 2}
    assert result["sheets"][1]["preview"] == ("city,count\nSeattle,1\n")


def test_read_file_bounds_tabular_schema_and_preview(
    tmp_path: Path, monkeypatch
) -> None:
    pytest.importorskip("polars")
    monkeypatch.chdir(tmp_path)

    headers = [f"column_{index}_{'x' * 40}" for index in range(120)]
    values = [f"value_{index}_{'y' * 40}" for index in range(120)]
    (tmp_path / "wide.csv").write_text(
        ",".join(headers) + "\n" + ",".join(values) + "\n",
        encoding="utf-8",
    )

    result = read_file_tool.handle({"path": "wide.csv"}, ToolContext())

    assert len(result["schema"]) <= read_file_tool.MAX_TABULAR_SCHEMA_CHARS
    assert len(result["preview"]) <= read_file_tool.MAX_TABULAR_PREVIEW_CHARS
    assert result["schema_truncated"] is True
    assert result["preview_truncated"] is True


def test_read_file_summarizes_pdf_content_and_metadata(
    tmp_path: Path, monkeypatch
) -> None:
    pytest.importorskip("pypdf")
    monkeypatch.chdir(tmp_path)

    from pypdf import PdfWriter

    pdf_path = tmp_path / "report.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.add_metadata({"/Title": "Quarterly Report", "/Author": "Agent"})
    with pdf_path.open("wb") as file_handle:
        writer.write(file_handle)

    result = read_file_tool.handle({"path": "report.pdf"}, ToolContext())

    assert result["metadata"]["pages"] == 1
    assert result["metadata"]["title"] == "Quarterly Report"
    assert result["metadata"]["author"] == "Agent"
    assert result["content"] == ""


def test_read_file_extracts_docx_text(tmp_path: Path, monkeypatch) -> None:
    pytest.importorskip("docx")
    monkeypatch.chdir(tmp_path)

    from docx import Document

    docx_path = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("Quarterly report")
    document.add_paragraph("Revenue grew 12%")
    document.save(docx_path)

    result = read_file_tool.handle({"path": "report.docx"}, ToolContext())

    assert result == {
        "path": "report.docx",
        "content": "Quarterly report\nRevenue grew 12%",
    }


def test_read_file_extracts_docx_tables(tmp_path: Path, monkeypatch) -> None:
    pytest.importorskip("docx")
    monkeypatch.chdir(tmp_path)

    from docx import Document

    docx_path = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("Summary")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "City"
    table.cell(0, 1).text = "Sales"
    table.cell(1, 0).text = "Seattle"
    table.cell(1, 1).text = "100"
    document.add_paragraph("End")
    document.save(docx_path)

    result = read_file_tool.handle({"path": "report.docx"}, ToolContext())

    assert result["content"] == "Summary\nCity\tSales\nSeattle\t100\nEnd"


def test_read_file_extracts_docx_headers_and_footers(
    tmp_path: Path, monkeypatch
) -> None:
    pytest.importorskip("docx")
    monkeypatch.chdir(tmp_path)

    from docx import Document

    docx_path = tmp_path / "report.docx"
    document = Document()
    section = document.sections[0]
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].text = "Company Inc."
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = "Page 1"
    document.add_paragraph("Body text")
    document.save(docx_path)

    result = read_file_tool.handle({"path": "report.docx"}, ToolContext())

    assert "Body text" in result["content"]
    assert "Company Inc." in result["content"]
    assert "Page 1" in result["content"]


def test_read_file_does_not_truncate_docx_content(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.chdir(tmp_path)
    (tmp_path / "report.docx").write_bytes(b"placeholder")
    full_text = "Quarterly report\n" + (
        "Revenue grew 12%. " * read_file_tool.MAX_READ_FILE_OUTPUT_CHARS
    )
    monkeypatch.setattr(read_file_tool, "_extract_docx_text", lambda path: full_text)

    result = read_file_tool.handle({"path": "report.docx"}, ToolContext())

    assert result == {
        "path": "report.docx",
        "content": full_text,
    }


def test_read_file_does_not_truncate_pdf_content(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.chdir(tmp_path)
    pdf_path = tmp_path / "report.pdf"
    pdf_path.write_bytes(b"%PDF-1.4")
    full_text = "Quarterly report\n" + (
        "Revenue grew 12%. " * read_file_tool.MAX_READ_FILE_OUTPUT_CHARS
    )

    fake_pypdf = ModuleType("pypdf")

    class FakePdfReader:
        def __init__(self, path: str) -> None:
            assert path == str(pdf_path)
            self.pages = [SimpleNamespace(extract_text=lambda: full_text)]
            self.metadata = SimpleNamespace(title="Quarterly Report", author="Agent")

    fake_pypdf.PdfReader = FakePdfReader
    monkeypatch.setitem(sys.modules, "pypdf", fake_pypdf)

    result = read_file_tool.handle({"path": "report.pdf"}, ToolContext())

    assert result == {
        "path": "report.pdf",
        "content": full_text,
        "metadata": {
            "pages": 1,
            "title": "Quarterly Report",
            "author": "Agent",
        },
    }


def test_read_file_reports_unreadable_docx_files(tmp_path: Path, monkeypatch) -> None:
    pytest.importorskip("docx")
    monkeypatch.chdir(tmp_path)
    (tmp_path / "broken.docx").write_bytes(b"not-a-valid-docx")

    result = read_file_tool.handle({"path": "broken.docx"}, ToolContext())

    assert result == {
        "error": "unable to read docx file: broken.docx is unreadable or corrupt"
    }


def test_read_file_allows_more_than_default_output_budget(
    tmp_path: Path, monkeypatch
) -> None:
    monkeypatch.chdir(tmp_path)
    long_line = (
        f"prefix-{'x' * (read_file_tool.MAX_READ_FILE_OUTPUT_CHARS // 2)}-suffix"
    )
    (tmp_path / "large.txt").write_text(f"{long_line}\n", encoding="utf-8")

    result = read_file_tool.handle({"path": "large.txt"}, ToolContext())

    assert "content_truncated" not in result
    assert result["content"] == f"{long_line}\n"


def test_read_file_bounds_very_large_content(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.chdir(tmp_path)
    long_line = (
        f"prefix-{'x' * (read_file_tool.MAX_READ_FILE_OUTPUT_CHARS + 200)}-suffix"
    )
    (tmp_path / "large.txt").write_text(f"{long_line}\n", encoding="utf-8")

    result = read_file_tool.handle({"path": "large.txt"}, ToolContext())

    assert result["content_truncated"] is True
    assert len(result["content"]) <= read_file_tool.MAX_READ_FILE_OUTPUT_CHARS
    assert result["content"].startswith("prefix-")
    assert result["content"].endswith("-suffix\n")
    assert "chars omitted" in result["content"]


def test_read_file_rejects_binary_files(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.chdir(tmp_path)
    (tmp_path / "blob.bin").write_bytes(b"\x00\x01\x02")

    result = read_file_tool.handle({"path": "blob.bin"}, ToolContext())

    assert "binary file is not supported" in result["error"]


def test_read_file_reports_empty_files_with_zero_range(
    tmp_path: Path, monkeypatch
) -> None:
    monkeypatch.chdir(tmp_path)
    (tmp_path / "empty.txt").write_text("", encoding="utf-8")

    result = read_file_tool.handle({"path": "empty.txt"}, ToolContext())

    assert result == {
        "path": "empty.txt",
        "start_line": 0,
        "end_line": 0,
        "total_lines": 0,
        "content": "",
        "has_more_lines": False,
        "empty": True,
    }
